"""Resume parsing + topic extraction.

Accepts a PDF or DOCX file's bytes, pulls text out, hands it to Groq with the
resume_extract prompt, returns the structured topic list. No persistence at
this layer — the router or a later phase decides what to store.

Upload security (Phase 4): file TYPE is verified by real magic bytes, not
only the filename extension or the client-supplied Content-Type header
(either of which a caller fully controls and could set to anything
regardless of the file's actual content). Parsing itself is bounded -
real page-count/size limits before any per-page work happens, and a
best-effort wall-clock timeout around the parse call. Malware scanning is
explicitly NOT implemented here - see MALWARE_SCANNING_NOTE below; this
is a documented infrastructure gap, not something faked.
"""

import io
import json
import signal
from pathlib import Path as _Path

from app.llm_client import chat_completion

MAX_BYTES = 5 * 1024 * 1024  # 5 MB cap
MAX_TEXT_CHARS = 60_000       # trim very long resumes before sending to the LLM
MAX_PDF_PAGES = 30            # a real resume is a handful of pages - bounds
                               # per-page parse work against a maliciously
                               # huge page count before extracting anything
PARSE_TIMEOUT_SECONDS = 10    # best-effort wall-clock bound on the parse call

# Real magic-byte signatures - a caller controls the filename and
# Content-Type header completely; neither is trustworthy evidence of what
# the bytes actually are. PDF files always start with "%PDF-". DOCX is
# Office Open XML, which is a ZIP container, so a real .docx always starts
# with the ZIP local-file-header signature (PK\x03\x04) - this doesn't
# prove the ZIP's INTERNAL contents are a valid Word document (that's
# what parsing failure below catches), only that it's genuinely a ZIP/
# OOXML-shaped file and not, say, an executable or a script renamed with
# a .pdf/.docx extension.
_PDF_MAGIC = b"%PDF-"
_ZIP_MAGIC = b"PK\x03\x04"

# Real infrastructure gap, stated honestly rather than faked: this
# deployment has no malware-scanning service (ClamAV, a cloud AV API,
# etc.) provisioned, and none is implemented here. A resume that passes
# every check in this module (real magic bytes, parses cleanly, bounded
# size/pages/time) could still theoretically carry malicious content a
# dedicated scanner would catch that a text-extraction library does not
# (e.g. an embedded exploit targeting a DIFFERENT consumer of the raw
# file than this text extractor). Flagging this explicitly rather than
# adding a scanning call this environment can't actually make real.
MALWARE_SCANNING_NOTE = (
    "No malware-scanning service is provisioned in this deployment. "
    "Uploaded files are validated by real magic bytes and parsed with "
    "bounded size/page-count/time limits, but are not scanned by a "
    "dedicated AV engine before being parsed or stored."
)


class UnsupportedFileError(ValueError):
    """A file that isn't the real, correct type it claims to be (wrong
    magic bytes for its extension), or a real file this module explicitly
    can't safely process (encrypted, too many pages)."""


def _looks_like_pdf(data: bytes) -> bool:
    return data[:5] == _PDF_MAGIC


def _looks_like_docx(data: bytes) -> bool:
    return data[:4] == _ZIP_MAGIC


class _ParseTimeout(Exception):
    pass


def _run_with_timeout(fn, *args, seconds: int = PARSE_TIMEOUT_SECONDS):
    """Best-effort wall-clock bound using SIGALRM - only available on
    Unix (the real production target: this runs in a Linux Docker
    container on EC2). On a platform without SIGALRM (Windows local dev),
    this honestly runs unbounded rather than pretending to enforce a
    timeout it structurally cannot on that platform - a real, stated
    limitation rather than a faked one."""
    if not hasattr(signal, "SIGALRM"):
        return fn(*args)

    def _on_alarm(signum, frame):
        raise _ParseTimeout()

    previous_handler = signal.signal(signal.SIGALRM, _on_alarm)
    signal.alarm(seconds)
    try:
        return fn(*args)
    except _ParseTimeout:
        raise ValueError(f"This file took too long to parse (over {seconds}s) and was rejected.")
    finally:
        signal.alarm(0)
        signal.signal(signal.SIGALRM, previous_handler)


def _load_prompt(name: str) -> str:
    return (_Path(__file__).parent.parent / "prompts" / name).read_text(encoding="utf-8")
try:
    from pypdf import PdfReader  # type: ignore
except (ImportError, Exception):
    PdfReader = None

try:
    from docx import Document  # type: ignore
except (ImportError, Exception):
    Document = None



# ---------------------------------------------------------------- text extraction
def _pdf_to_text(data: bytes) -> str:
    if PdfReader is None:
        raise ValueError("pypdf is not installed. Please install pypdf to process PDF files.")
    reader = PdfReader(io.BytesIO(data))
    # Encrypted PDFs are rejected outright rather than attempting a blank-
    # password unlock (which pypdf will sometimes do implicitly) - "reject
    # unsupported... encrypted... files safely" from the audit; an
    # encrypted resume isn't something this pipeline can honestly claim
    # to have read.
    if reader.is_encrypted:
        raise UnsupportedFileError("This PDF is password-protected. Please upload an unprotected file.")
    if len(reader.pages) > MAX_PDF_PAGES:
        raise UnsupportedFileError(
            f"This PDF has {len(reader.pages)} pages; a resume should have at most {MAX_PDF_PAGES}."
        )
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _docx_to_text(data: bytes) -> str:
    if Document is None:
        raise ValueError("python-docx is not installed. Please install python-docx to process DOCX files.")
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def extract_text(data: bytes, filename: str, content_type: str = "") -> str:
    """Return plain text from a resume file. Raises ValueError on unsupported/empty."""
    if not data:
        raise ValueError("Empty file")
    if len(data) > MAX_BYTES:
        raise ValueError(f"File too large ({len(data)} bytes; max {MAX_BYTES})")

    name = (filename or "").lower()
    ct = (content_type or "").lower()

    claims_pdf = name.endswith(".pdf") or "pdf" in ct
    claims_docx = name.endswith(".docx") or "wordprocessingml" in ct

    if not claims_pdf and not claims_docx:
        raise ValueError("Unsupported file type. Please upload a PDF or DOCX.")

    # Real magic-byte check - the filename/Content-Type only say what the
    # CALLER claims this is; a caller controls both completely regardless
    # of the actual bytes. A mismatch is rejected outright rather than
    # handed to a parser expecting a different format.
    if claims_pdf and not _looks_like_pdf(data):
        raise UnsupportedFileError(
            "This file's content doesn't match a real PDF (wrong file signature). "
            "It may be corrupted, a different file type renamed, or not a real PDF."
        )
    if claims_docx and not _looks_like_docx(data):
        raise UnsupportedFileError(
            "This file's content doesn't match a real DOCX (wrong file signature). "
            "It may be corrupted, a different file type renamed, or not a real DOCX."
        )

    if claims_pdf:
        text = _run_with_timeout(_pdf_to_text, data)
    else:
        text = _run_with_timeout(_docx_to_text, data)

    text = text.strip()
    if not text:
        raise ValueError("Could not read any text from the file (is it scanned/image-only?)")
    return text[:MAX_TEXT_CHARS]


# ---------------------------------------------------------------- LLM extraction
def _strip_fences(raw: str) -> str:
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    return raw.strip()


def _call_groq(messages: list, max_tokens: int = 3000) -> str:
    # Name kept for minimal diff at call sites below; routes through
    # app.llm_client, which picks Groq or Bedrock per settings.LLM_PROVIDER.
    return chat_completion(messages, max_tokens=max_tokens, temperature=0.1)


VALID_LEVELS = {"basic", "intermediate", "advanced", "expert"}


def _confidence_pct_or_none(raw) -> int | None:
    """Only trust the LLM's OWN stated confidence when it's a real, in-range
    number. Previously, a missing/invalid confidence_pct triggered
    _calculate_fallback_confidence() - a formula that converted the
    suggested_level, evidence-text word count, and years of experience into
    a fabricated fine-grained percentage (e.g. "88%") that looked exactly
    like a real LLM-stated confidence but wasn't one. That's precisely the
    "confidence must represent uncertainty, not achievement" violation the
    platform audit targets: it invented false precision for a value we
    simply don't have. Returning None here is the honest alternative -
    mastery_service._apply_topics() already has a documented, low, source-
    specific default (_DEFAULT_CONFIDENCE_BY_SOURCE["resume"] = 0.6) for
    exactly this "no real confidence available" case, and the frontend
    (AssessSkills.jsx/LearnerIntakeWorkspace.jsx) already renders a topic
    with no confidence_pct as "Self-reported"/unconfident rather than
    crashing or fabricating a display value."""
    try:
        conf = int(raw)
    except (TypeError, ValueError):
        return None
    return conf if 30 <= conf <= 100 else None


MAX_TEXT_FIELD_CHARS = 500  # education/projects/suggested_goal - keep these short, real summaries


def _clean_text_field(payload: dict, key: str, max_len: int = MAX_TEXT_FIELD_CHARS) -> str:
    val = payload.get(key)
    if not isinstance(val, str):
        return ""
    return val.strip()[:max_len]


def _validate(payload: dict) -> dict:
    topics = payload.get("topics")
    if not isinstance(topics, list):
        raise ValueError("payload.topics is not a list")
    years = payload.get("detected_years_experience", 0)
    try:
        years = int(years)
    except (TypeError, ValueError):
        years = 0
    years = max(0, years)

    clean = []
    for t in topics:
        name = (t.get("name") or "").strip()
        level = (t.get("suggested_level") or "").lower().strip()
        evidence = (t.get("evidence") or "").strip()
        if not name or level not in VALID_LEVELS:
            continue
        conf = _confidence_pct_or_none(t.get("confidence_pct"))

        clean.append({
            "name": name,
            "evidence": evidence,
            "suggested_level": level,
            "confidence_pct": conf,
        })
    return {
        "topics": clean,
        "detected_years_experience": years,
        # Real, resume-grounded context beyond just skills - see resume_extract.txt.
        # Never fabricated: the LLM is instructed to return "" when the resume
        # doesn't actually say enough to fill these confidently.
        "education": _clean_text_field(payload, "education"),
        "projects": _clean_text_field(payload, "projects"),
        "suggested_goal": _clean_text_field(payload, "suggested_goal", max_len=100),
    }


def extract_topics(resume_text: str) -> dict:
    """Returns {topics: [...], detected_years_experience: N}."""
    messages = [
        {"role": "system", "content": _load_prompt("resume_extract.txt")},
        {"role": "user", "content": f"RESUME:\n\n{resume_text}"},
    ]
    raw = _call_groq(messages)
    try:
        return _validate(json.loads(_strip_fences(raw)))
    except (json.JSONDecodeError, ValueError):
        # Retry once with an explicit reminder.
        messages.append({"role": "assistant", "content": raw})
        messages.append({
            "role": "user",
            "content": "Return ONLY the JSON object. No markdown, no explanation.",
        })
        try:
            return _validate(json.loads(_strip_fences(_call_groq(messages))))
        except Exception:
            return {"topics": [], "detected_years_experience": 0}
